"""
core.py - document anonymization logic, independent from the GUI.
"""
from __future__ import annotations

import csv
import json
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Callable, Optional


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class Span:
    file: str
    label: str
    text: str
    start: int
    end: int
    score: Optional[float]
    will_redact: bool = True
    code: str = ""   # e.g. "PERSON_1" - assigned by build_entity_registry
    source: str = "opf"  # opf, regex, manual, merged


@dataclass
class ProcessingResult:
    spans: list[Span] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)


SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".xlsx", ".txt", ".csv", ".tsv")
SAFE_OUTPUT_SUBDIR = "01_LLM_SAFE_FILES"
RESERVED_OUTPUT_SUBDIR = "99_RESERVED_DO_NOT_UPLOAD"


# ---------------------------------------------------------------------------
# File extraction
# ---------------------------------------------------------------------------

def _read_pdf(path: str) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency: install pdfplumber with 'py -m pip install pdfplumber'."
        ) from exc

    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text


def _table_to_text(table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            cell_parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
            for nested in cell.tables:
                nested_text = _table_to_text(nested).strip()
                if nested_text:
                    cell_parts.append(nested_text)
            cells.append(" ".join(cell_parts))
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _read_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency: install python-docx with 'py -m pip install python-docx'."
        ) from exc

    doc = Document(path)
    parts: list[str] = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        table_text = _table_to_text(table).strip()
        if table_text:
            parts.append(table_text)

    # Headers and footers often contain customer codes, names, or references.
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in container.tables:
                table_text = _table_to_text(table).strip()
                if table_text:
                    parts.append(table_text)

    return "\n".join(parts)


def _format_cell_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "YES" if value else "NO"
    return str(value).strip()


_XLSX_CONTEXT_HEADERS = re.compile(
    r"\b(supplier|vendor|customer|client|company|organization|entity|business\s+name|"
    r"fornitore|cliente|ragione\s+sociale|azienda|ditta|societ[aà]|"
    r"contact|name|first\s+name|last\s+name|nome|cognome|"
    r"email|mail|phone|mobile|telefono|cellulare|address|indirizzo|"
    r"vat|tax\s+id|fiscal\s+id|p\.?\s*iva|partita\s+iva|codice\s+fiscale|iban)\b",
    re.IGNORECASE,
)


def _looks_like_header(values: list[str]) -> bool:
    filled = [value for value in values if value]
    if len(filled) < 2:
        return False
    textish = sum(1 for value in filled if re.search(r"[A-Za-z]", value))
    return textish >= max(2, len(filled) // 2)


def _format_xlsx_row(values: list[str], headers: list[str] | None, row_number: int) -> str:
    if not headers:
        return " | ".join(values)

    parts: list[str] = []
    for idx, value in enumerate(values):
        if not value:
            continue
        header = headers[idx] if idx < len(headers) else ""
        if header and _XLSX_CONTEXT_HEADERS.search(header):
            parts.append(f"{header}: {value}")
        else:
            parts.append(value)
    if not parts:
        return ""
    return f"Row {row_number}: " + " | ".join(parts)


def _read_xlsx(path: str) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency: install openpyxl with 'py -m pip install openpyxl'."
        ) from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets:
            rows: list[str] = []
            headers: list[str] | None = None
            for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = [_format_cell_value(value) for value in row]
                while values and not values[-1]:
                    values.pop()
                if not any(values):
                    continue
                if headers is None and _looks_like_header(values):
                    headers = values
                    rows.append(" | ".join(values))
                    continue
                formatted = _format_xlsx_row(values, headers, row_number)
                if formatted:
                    rows.append(formatted)
            if rows:
                parts.append(f"[Sheet: {sheet.title}]")
                parts.extend(rows)
    finally:
        workbook.close()

    return "\n".join(parts)


def _read_delimited(path: str, delimiter: str) -> str:
    text = _read_text(path)
    rows = list(csv.reader(text.splitlines(), delimiter=delimiter))
    if not rows:
        return ""

    headers = [_format_cell_value(value) for value in rows[0]]
    has_header = _looks_like_header(headers)
    output: list[str] = []
    if has_header:
        output.append(" | ".join(headers))
        data_rows = rows[1:]
    else:
        data_rows = rows
        headers = []

    for row_number, row in enumerate(data_rows, start=2 if has_header else 1):
        values = [_format_cell_value(value) for value in row]
        while values and not values[-1]:
            values.pop()
        if not any(values):
            continue
        formatted = _format_xlsx_row(values, headers if has_header else None, row_number)
        if formatted:
            output.append(formatted)
    return "\n".join(output)


def _read_text(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "utf-16", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeError:
            continue
    raise ValueError(f"Could not read {path} with any known encoding")


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".xlsx":
        return _read_xlsx(path)
    if ext == ".csv":
        return _read_delimited(path, ",")
    if ext == ".tsv":
        return _read_delimited(path, "\t")
    if ext == ".txt":
        return _read_text(path)
    raise ValueError(f"Unsupported format: {ext}")


# ---------------------------------------------------------------------------
# PII detection
# ---------------------------------------------------------------------------

def _digits(value: str) -> str:
    return re.sub(r"\D", "", value)


def _valid_piva(value: str) -> bool:
    digits = _digits(value)
    if len(digits) != 11:
        return False
    total = sum(int(digits[i]) for i in range(0, 10, 2))
    for i in range(1, 10, 2):
        doubled = int(digits[i]) * 2
        total += doubled if doubled < 10 else doubled - 9
    check = (10 - total % 10) % 10
    return check == int(digits[-1])


_CF_ODD = {
    "0": 1, "1": 0, "2": 5, "3": 7, "4": 9, "5": 13, "6": 15, "7": 17,
    "8": 19, "9": 21, "A": 1, "B": 0, "C": 5, "D": 7, "E": 9, "F": 13,
    "G": 15, "H": 17, "I": 19, "J": 21, "K": 2, "L": 4, "M": 18, "N": 20,
    "O": 11, "P": 3, "Q": 6, "R": 8, "S": 12, "T": 14, "U": 16, "V": 10,
    "W": 22, "X": 25, "Y": 24, "Z": 23,
}
_CF_EVEN = {str(i): i for i in range(10)} | {chr(ord("A") + i): i for i in range(26)}


def _valid_italian_tax_code(value: str) -> bool:
    cf = re.sub(r"\s", "", value).upper()
    if not re.fullmatch(r"[A-Z]{6}\d{2}[A-EHLMPRST]\d{2}[A-Z]\d{3}[A-Z]", cf):
        return False
    total = 0
    for idx, char in enumerate(cf[:15], start=1):
        total += _CF_ODD[char] if idx % 2 else _CF_EVEN[char]
    expected = chr(ord("A") + (total % 26))
    return expected == cf[-1]


def _valid_iban(value: str) -> bool:
    iban = re.sub(r"\s", "", value).upper()
    if not re.fullmatch(r"IT\d{2}[A-Z]\d{10}[A-Z0-9]{12}", iban):
        return False
    rearranged = iban[4:] + iban[:4]
    numeric = "".join(str(ord(ch) - 55) if ch.isalpha() else ch for ch in rearranged)
    remainder = 0
    for ch in numeric:
        remainder = (remainder * 10 + int(ch)) % 97
    return remainder == 1


def _valid_phone(value: str) -> bool:
    digits = _digits(value)
    if value.strip().startswith("+39"):
        digits = digits[2:] if digits.startswith("39") else digits
    return 8 <= len(digits) <= 11 and digits[0] in {"0", "3"}


@dataclass(frozen=True)
class RegexRule:
    label: str
    pattern: re.Pattern
    group: int = 0
    validator: Optional[Callable[[str], bool]] = None


_REGEX_RULES: list[RegexRule] = [
    RegexRule(
        "private_email",
        re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE),
    ),
    RegexRule(
        "private_id",
        re.compile(r"\b(?:P\.?\s*IVA\s*:?\s*)?(\d{11})\b", re.IGNORECASE),
        group=1,
        validator=_valid_piva,
    ),
    RegexRule(
        "private_id",
        re.compile(r"\b[A-Z]{6}\d{2}[A-EHLMPRST]\d{2}[A-Z]\d{3}[A-Z]\b", re.IGNORECASE),
        validator=_valid_italian_tax_code,
    ),
    RegexRule(
        "private_id",
        re.compile(r"\bIT\s*\d{2}\s*[A-Z]\s*(?:\d\s*){10}(?:[A-Z0-9]\s*){12}\b", re.IGNORECASE),
        validator=_valid_iban,
    ),
    RegexRule(
        "private_phone",
        re.compile(
            r"(?<!\w)(?:\+39\s*)?(?:0\d{1,4}|3\d{2})[\s./-]?\d{3,4}[\s./-]?\d{3,4}(?!\w)",
            re.IGNORECASE,
        ),
        validator=_valid_phone,
    ),
]

_REGEX_LABEL_PRIORITY = {
    "private_id": 5,
    "private_email": 4,
    "private_phone": 3,
    "private_address": 2,
}

_CONTEXT_ORG_RE = re.compile(
    r"(?im)\b("
    r"supplier|vendor|customer|client|company|organization|entity|business\s+name|"
    r"fornitore|ragione\s+sociale|azienda|ditta|societ[aà]|cliente"
    r")\s*[:=\-]\s*(?P<value>[^\n\r|;]{2,120})"
)

# Italian street prefixes used to isolate the street-level part of an address.
_STREET_RE = re.compile(
    r"\b(Via|Viale|Vicolo|Corso|Piazza|Largo|Strada|Contrada|Borgata|Borgo|"
    r"Localita|Località|Loc\.?|Frazione|Fraz\.?|V\.le|C\.so|P\.za|Sp(?:\.|\b)|"
    r"S\.?S\.?|Statale|Provinciale)\b"
    r"[\s\S]{1,80}?\d+\s*[A-Za-z]?(?:\s*/\s*[A-Za-z\d]+)?",
    re.IGNORECASE,
)


def _normalize_free_text(value: str) -> str:
    return " ".join(value.split()).casefold()


def normalize_entity_text(label: str, text: str) -> str:
    value = " ".join(text.split()).strip()
    if label == "private_email":
        return value.casefold()
    if label == "private_phone":
        digits = _digits(value)
        if digits.startswith("39") and len(digits) > 10:
            digits = digits[2:]
        return digits
    if label == "private_id":
        return re.sub(r"[\s./-]", "", value).upper()
    return value.casefold()


_OPF_LABEL_ALIASES = {
    "account_number": "private_id",
    "bank_account": "private_id",
    "credit_card": "private_id",
    "fiscal_id": "private_id",
    "iban": "private_id",
    "id_number": "private_id",
    "tax_id": "private_id",
    "vat_number": "private_id",
    "private_account_number": "private_id",
    "private_bank_account": "private_id",
    "private_credit_card": "private_id",
    "private_fiscal_id": "private_id",
    "private_iban": "private_id",
    "private_tax_id": "private_id",
    "private_vat_number": "private_id",
}


def canonical_label(label: str) -> str:
    return _OPF_LABEL_ALIASES.get(label, label)


def _is_excluded(span: Span, exclude_terms: set[str]) -> bool:
    if not exclude_terms:
        return False
    return (
        _normalize_free_text(span.text) in exclude_terms
        or normalize_entity_text(span.label, span.text) in exclude_terms
    )


def _refine_address_span(span: Span) -> Span:
    """
    Reduce a private_address span to the street-level part where possible.
    If no recognizable street prefix is found, keep the original span.
    """
    m = _STREET_RE.search(span.text)
    if not m:
        return span
    street_start = span.start + m.start()
    street_end = span.start + m.end()
    street_text = span.text[m.start():m.end()].strip()
    return Span(
        file=span.file,
        label=span.label,
        text=street_text,
        start=street_start,
        end=street_end,
        score=span.score,
        will_redact=span.will_redact,
        code=span.code,
        source=span.source,
    )


def _ranges_overlap(start_a: int, end_a: int, start_b: int, end_b: int) -> bool:
    return max(start_a, start_b) < min(end_a, end_b)


def _detect_regex_spans(
    text: str,
    basename: str,
    labels_filter: set[str],
    exclude_terms: set[str],
) -> list[Span]:
    """Detect PII with regex rules, leaving overlap resolution to detect_spans."""
    spans: list[Span] = []
    for rule in _REGEX_RULES:
        for m in rule.pattern.finditer(text):
            candidate = m.group(rule.group).strip()
            if rule.validator and not rule.validator(candidate):
                continue
            start = m.start(rule.group)
            end = m.end(rule.group)
            span = Span(
                file=basename,
                label=rule.label,
                text=candidate,
                start=start,
                end=end,
                score=None,
                will_redact=(rule.label in labels_filter),
                source="regex",
            )
            if _is_excluded(span, exclude_terms):
                span.will_redact = False
            spans.append(span)

    accepted: list[Span] = []
    for span in sorted(
        spans,
        key=lambda s: (
            s.start,
            -_REGEX_LABEL_PRIORITY.get(s.label, 0),
            -(s.end - s.start),
        ),
    ):
        if any(_ranges_overlap(span.start, span.end, other.start, other.end) for other in accepted):
            continue
        accepted.append(span)
    return sorted(accepted, key=lambda s: (s.start, s.end, s.label))


def _clean_context_value(value: str) -> str:
    value = value.strip(" \t:-=;|")
    value = re.split(
        r"\s+(?:VAT(?:\s+number)?|Tax\s+ID|Fiscal\s+ID|P\.?\s*IVA|Partita\s+IVA|"
        r"Codice\s+fiscale|CF|Email|Mail|Phone|Mobile|Telefono|Tel\.?|IBAN|Address|Indirizzo)\s*[:=\-]",
        value,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0]
    return value.strip(" \t:-=;|")


def _detect_context_spans(
    text: str,
    basename: str,
    labels_filter: set[str],
    exclude_terms: set[str],
) -> list[Span]:
    """Detect values after supplier/customer fields such as Fornitore or Ragione sociale."""
    spans: list[Span] = []
    for match in _CONTEXT_ORG_RE.finditer(text):
        raw = match.group("value")
        cleaned = _clean_context_value(raw)
        if len(cleaned) < 2:
            continue
        raw_start = match.start("value")
        leading = len(raw) - len(raw.lstrip(" \t:-=;|"))
        start = raw_start + leading
        end = start + len(cleaned)
        span = Span(
            file=basename,
            label="private_organization",
            text=cleaned,
            start=start,
            end=end,
            score=None,
            will_redact=("private_organization" in labels_filter),
            source="context",
        )
        if _is_excluded(span, exclude_terms):
            span.will_redact = False
        spans.append(span)
    return dedupe_spans(spans)


def _load_model():
    from opf import OPF  # lazy import: keep startup usable when OPF is missing
    return OPF(device="cpu", output_mode="typed")


def _source_priority(source: str) -> int:
    return {"manual": 4, "regex": 3, "context": 3, "opf": 2, "merged": 1}.get(source, 0)


def _span_rank(span: Span) -> tuple[int, float, int]:
    return (
        _source_priority(span.source),
        span.score if span.score is not None else 0.0,
        span.end - span.start,
    )


def _trim_detected_bounds(text: str, start: int, end: int) -> tuple[int, int, str]:
    while start < end and text[start] in " \t\r\n,;|":
        start += 1
    while end > start and text[end - 1] in " \t\r\n,;|":
        end -= 1
    return start, end, text[start:end].strip()


def _append_with_overlap_priority(spans: list[Span], candidate: Span) -> None:
    overlaps = [s for s in spans if _ranges_overlap(candidate.start, candidate.end, s.start, s.end)]
    if not overlaps:
        spans.append(candidate)
        return
    if all(_span_rank(candidate) > _span_rank(existing) for existing in overlaps):
        for existing in overlaps:
            spans.remove(existing)
        spans.append(candidate)


def dedupe_spans(spans: list[Span]) -> list[Span]:
    """Remove exact duplicates while keeping the most reliable span."""
    best_by_key: dict[tuple, Span] = {}
    for span in spans:
        if span.start < 0 or span.end <= span.start:
            continue
        key = (
            span.file,
            span.label,
            span.start,
            span.end,
            normalize_entity_text(span.label, span.text),
        )
        current = best_by_key.get(key)
        if current is None or _span_rank(span) > _span_rank(current):
            best_by_key[key] = span
    return sorted(best_by_key.values(), key=lambda s: (s.file, s.start, s.end, s.label))


def detect_spans(
    files: list[str],
    labels_filter: set[str],
    progress_cb: Optional[Callable[[int, int, str], None]] = None,
    min_score: Optional[float] = None,
    exclude_terms: Optional[set[str]] = None,
) -> ProcessingResult:
    """
    Analyze files and return detected spans.
    progress_cb(current, total, filename) is called for each file.
    """
    model = _load_model()
    result = ProcessingResult()
    total = len(files)
    exclusions = {_normalize_free_text(term) for term in (exclude_terms or set()) if term.strip()}

    for i, filepath in enumerate(files):
        basename = os.path.basename(filepath)
        if progress_cb:
            progress_cb(i, total, basename)
        try:
            text = extract_text(filepath)
            detected = model.redact(text)
            opf_spans: list[Span] = []
            for s in detected.detected_spans:
                score = getattr(s, "score", None)
                if min_score is not None and score is not None and score < min_score:
                    continue
                start, end, detected_text = _trim_detected_bounds(text, s.start, s.end)
                if not detected_text:
                    continue
                label = canonical_label(s.label)
                span = Span(
                    file=basename,
                    label=label,
                    text=detected_text,
                    start=start,
                    end=end,
                    score=score,
                    will_redact=(label in labels_filter),
                    source="opf",
                )
                if label == "private_address":
                    span = _refine_address_span(span)
                if _is_excluded(span, exclusions):
                    span.will_redact = False
                opf_spans.append(span)

            # Add context/regex spans, replacing lower-priority model spans where needed.
            context_spans = _detect_context_spans(text, basename, labels_filter, exclusions)
            for cs in context_spans:
                _append_with_overlap_priority(opf_spans, cs)

            regex_spans = _detect_regex_spans(text, basename, labels_filter, exclusions)
            for rs in regex_spans:
                _append_with_overlap_priority(opf_spans, rs)

            result.spans.extend(dedupe_spans(opf_spans))
        except Exception as exc:
            result.errors.append(f"{basename}: {exc}")

    if progress_cb:
        progress_cb(total, total, "")
    return result


# ---------------------------------------------------------------------------
# Entity registry - assigns stable progressive codes across all files
# ---------------------------------------------------------------------------

def build_entity_registry(spans: list[Span], redacted_only: bool = False) -> dict[tuple, dict]:
    """
    Assign a progressive code per normalized entity and label.
    """
    registry: dict[tuple, dict] = {}
    counters: dict[str, int] = {}
    ordered_spans = sorted(spans, key=lambda s: (s.file, s.start, s.end))
    for s in ordered_spans:
        if redacted_only and not s.will_redact:
            continue
        normalized = normalize_entity_text(s.label, s.text)
        key = (s.label, normalized)
        if key not in registry:
            prefix = LABEL_PREFIX.get(s.label, s.label.upper())
            counters[prefix] = counters.get(prefix, 0) + 1
            code = f"{prefix}_{counters[prefix]}"
            registry[key] = {
                "entity_id": len(registry) + 1,
                "placeholder": code,
                "label": s.label,
                "text": s.text,
            }
        s.code = registry[key]["placeholder"]
    return registry


# ---------------------------------------------------------------------------
# Redaction
# ---------------------------------------------------------------------------

def _merge_redaction_group(text: str, group: list[Span]) -> Span:
    start = min(s.start for s in group)
    end = max(s.end for s in group)
    best = max(group, key=_span_rank)
    if len(group) == 1 and best.start == start and best.end == end:
        return best
    return Span(
        file=best.file,
        label=best.label,
        text=text[start:end],
        start=start,
        end=end,
        score=best.score,
        will_redact=True,
        code=best.code,
        source="merged",
    )


def _resolve_redaction_spans(text: str, spans: list[Span]) -> list[Span]:
    candidates = sorted(
        [s for s in spans if s.will_redact and s.start >= 0 and s.end > s.start],
        key=lambda s: (s.start, s.end),
    )
    if not candidates:
        return []

    resolved: list[Span] = []
    group: list[Span] = [candidates[0]]
    group_end = candidates[0].end

    for span in candidates[1:]:
        if span.start < group_end:
            group.append(span)
            group_end = max(group_end, span.end)
        else:
            resolved.append(_merge_redaction_group(text, group))
            group = [span]
            group_end = span.end
    resolved.append(_merge_redaction_group(text, group))
    return resolved


def _redact_text(text: str, spans: list[Span]) -> str:
    for s in sorted(_resolve_redaction_spans(text, spans), key=lambda x: x.start, reverse=True):
        placeholder = f"[{s.code}]" if s.code else f"[{s.label.upper()}]"
        text = text[:s.start] + placeholder + text[s.end:]
    return text


def preview_redacted_text(text: str, spans: list[Span]) -> str:
    """Return redacted text for GUI preview."""
    return _redact_text(text, spans)


def safe_output_dir(output_dir: str) -> str:
    return os.path.join(output_dir, SAFE_OUTPUT_SUBDIR)


def reserved_output_dir(output_dir: str) -> str:
    return os.path.join(output_dir, RESERVED_OUTPUT_SUBDIR)


def _output_path_for(input_path: str, output_dir: str) -> str:
    basename = os.path.basename(input_path)
    stem, ext = os.path.splitext(basename)
    ext_tag = ext.lower().lstrip(".") or "file"
    return os.path.join(output_dir, f"{stem}_{ext_tag}_anonymized.txt")


def save_outputs(
    files: list[str],
    all_spans: list[Span],
    output_dir: str,
    progress_cb: Optional[Callable[[int, int, str], None]] = None,
    include_kept_in_index: bool = False,
) -> list[str]:
    """
    Save anonymized files and CSV/JSON mapping files.
    Return a list of errors, if any.
    """
    os.makedirs(output_dir, exist_ok=True)
    safe_dir = safe_output_dir(output_dir)
    reserved_dir = reserved_output_dir(output_dir)
    os.makedirs(safe_dir, exist_ok=True)
    os.makedirs(reserved_dir, exist_ok=True)
    errors: list[str] = []
    total = len(files)

    # Sensitive mapping files include only actually redacted spans by default.
    registry = build_entity_registry(all_spans, redacted_only=True)

    spans_by_file: dict[str, list[Span]] = {}
    for s in all_spans:
        spans_by_file.setdefault(s.file, []).append(s)

    for i, filepath in enumerate(files):
        basename = os.path.basename(filepath)
        if progress_cb:
            progress_cb(i, total, basename)
        try:
            text = extract_text(filepath)
            file_spans = spans_by_file.get(basename, [])
            redacted = _redact_text(text, file_spans)
            out_path = _output_path_for(filepath, safe_dir)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(redacted)
        except Exception as exc:
            errors.append(f"{basename}: {exc}")

    if progress_cb:
        progress_cb(total, total, "")

    _save_mapping(registry, reserved_dir)
    _save_index(all_spans, reserved_dir, include_kept=include_kept_in_index)
    _save_log(files, all_spans, errors, reserved_dir)
    return errors


def _save_mapping(registry: dict[tuple, dict], output_dir: str) -> None:
    rows = list(registry.values())

    csv_path = os.path.join(output_dir, "mapping_entities.csv")
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["entity_id", "placeholder", "label", "text"])
        writer.writeheader()
        writer.writerows(rows)

    json_path = os.path.join(output_dir, "mapping_entities.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(rows, f, ensure_ascii=False, indent=2)


def _save_index(spans: list[Span], output_dir: str, include_kept: bool = False) -> None:
    csv_path = os.path.join(output_dir, "index_occurrences.csv")
    fields = ["file", "placeholder", "label", "text", "start", "end", "score", "will_redact", "source"]
    rows = spans if include_kept else [s for s in spans if s.will_redact]
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for s in rows:
            writer.writerow({
                "file": s.file,
                "placeholder": s.code,
                "label": s.label,
                "text": s.text,
                "start": s.start,
                "end": s.end,
                "score": f"{s.score:.3f}" if s.score is not None else "",
                "will_redact": s.will_redact,
                "source": s.source,
            })


def _save_log(files: list[str], spans: list[Span], errors: list[str], output_dir: str) -> None:
    redacted = [s for s in spans if s.will_redact]
    kept = [s for s in spans if not s.will_redact]
    by_label: dict[str, int] = {}
    for s in redacted:
        by_label[s.label] = by_label.get(s.label, 0) + 1

    lines = [
        f"Processed at: {datetime.now().isoformat(timespec='seconds')}",
        f"Processed files: {len(files)}",
        f"Total spans: {len(spans)}",
        f"Redacted spans: {len(redacted)}",
        f"Kept spans: {len(kept)}",
        "",
        "Input files:",
        *[f"- {os.path.basename(path)}" for path in files],
        "",
        "Redactions by label:",
        *[f"- {label}: {count}" for label, count in sorted(by_label.items())],
        "",
    ]
    if errors:
        lines.append("Errors:")
        lines.extend(f"- {err}" for err in errors)
    else:
        lines.append("No errors.")

    with open(os.path.join(output_dir, "processing_log.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def collect_files(folder: str) -> list[str]:
    files: list[str] = []
    for entry in os.scandir(folder):
        if not entry.is_file():
            continue
        if os.path.splitext(entry.name)[1].lower() in SUPPORTED_EXTENSIONS:
            files.append(entry.path)
    return sorted(files, key=lambda p: os.path.basename(p).casefold())


def collect_unsupported_files(folder: str) -> list[str]:
    unsupported: list[str] = []
    for entry in os.scandir(folder):
        if not entry.is_file():
            continue
        ext = os.path.splitext(entry.name)[1].lower()
        if ext and ext not in SUPPORTED_EXTENSIONS:
            unsupported.append(entry.path)
    return sorted(unsupported, key=lambda p: os.path.basename(p).casefold())


DEFAULT_LABELS = {
    "private_person",
    "private_email",
    "private_phone",
    "private_address",
    "private_organization",
    "private_id",
}

ALL_KNOWN_LABELS = DEFAULT_LABELS | {
    "private_date",
    "private_location",
}

# Human-readable placeholder prefix for each label.
LABEL_PREFIX: dict[str, str] = {
    "private_person":       "PERSON",
    "private_organization": "ORGANIZATION",
    "private_email":        "EMAIL",
    "private_phone":        "PHONE",
    "private_address":      "ADDRESS",
    "private_id":           "ID",
    "private_date":         "DATE",
    "private_location":     "LOCATION",
}
